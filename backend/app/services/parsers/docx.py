import io
import docx

def parse_docx(file_bytes: bytes) -> tuple[str, float, str]:
    """
    Parses a DOCX file preserving structure (headings, tables, lists, paragraph groupings).
    Returns:
        extracted_text (str)
        confidence (float): Always 100.0
        warning_message (str): Always empty
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    output_parts = []
    
    # Process elements in order
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            # Find the paragraph object corresponding to the element
            p = None
            for doc_p in doc.paragraphs:
                if doc_p._element == element:
                    p = doc_p
                    break
            
            if p:
                text = p.text.strip()
                if not text:
                    continue
                
                # Check heading style
                style_name = p.style.name.lower()
                if 'heading' in style_name:
                    # Determine level
                    level = 1
                    for char in style_name:
                        if char.isdigit():
                            level = int(char)
                            break
                    output_parts.append(f"\n{'#' * level} {text}\n")
                elif style_name.startswith('list'):
                    output_parts.append(f"- {text}")
                else:
                    output_parts.append(text)
                    
        elif element.tag.endswith('tbl'):  # Table
            table = None
            for doc_t in doc.tables:
                if doc_t._element == element:
                    table = doc_t
                    break
            
            if table:
                table_md = []
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_md.append("| " + " | ".join(row_data) + " |")
                    
                    # Add divider after header row
                    if row_idx == 0:
                        table_md.append("|" + "|".join(["---"] * len(row_data)) + "|")
                        
                output_parts.append("\n" + "\n".join(table_md) + "\n")
                
    full_text = "\n\n".join(output_parts).strip()
    return full_text, 100.0, ""
